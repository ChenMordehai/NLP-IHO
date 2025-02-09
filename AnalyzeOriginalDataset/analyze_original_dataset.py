"""
This script provides a graphical user interface (GUI) for analyzing text data within datasets and generating a report in DOCX format. 

Features:
- GUI for easy interaction and data input using tkinter.
- Supports loading of data from both Excel (.xlsx) and CSV (.csv) files.
- Automatic detection of required data columns and verification of their presence.
- Calculation of basic statistics such as text length, date ranges, and language distribution.
- Generation of graphical data distributions, including histograms for text length and language usage.
- Creation of a detailed DOCX report that includes dataset overview, text content details, language analysis, and graphical distributions.
- Language detection to categorize texts based on the proportion of English and Hebrew content, using the lingua library.
- Error handling to manage file loading/saving issues and data inconsistencies.

Usage:
- Run the script to open the GUI.
- Fill in the required fields, browse and select the input data file and output folder.
- Submit the form to process the data and generate the report. 
"""
import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
import matplotlib.pyplot as plt
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import os
import datetime
from lingua import Language, LanguageDetectorBuilder

languages = [Language.ENGLISH, Language.HEBREW]
detector = LanguageDetectorBuilder.from_languages(*languages).build()


def calculate_language_ratios(df, text_column='Text'):
    """
    Calculate the Hebrew-English ratio for each row in the text column, considering the proportion of words.
    """
    hebrew_count = 0
    english_count = 0
    mixed_count = 0

    for text in df[text_column]:
        if pd.isna(text):  # Skip empty texts
            continue

        detected_languages = detector.detect_multiple_languages_of(text)
        hebrew_words = 0
        english_words = 0
        total_words = len(text.split())

        if detected_languages:
            for segment in detected_languages:
                language = segment.language
                segment_word_count = segment.word_count

                if language == Language.HEBREW:
                    hebrew_words += segment_word_count
                elif language == Language.ENGLISH:
                    english_words += segment_word_count

            if (english_words > 0.1 * total_words) and (hebrew_words > 0.1 * total_words):  # threshold for mixed
                mixed_count += 1
            elif hebrew_words > english_words:
                hebrew_count += 1
            else:
                english_count += 1

    total_texts = hebrew_count + english_count + mixed_count
    hebrew_ratio = hebrew_count / total_texts if total_texts > 0 else 0
    english_ratio = english_count / total_texts if total_texts > 0 else 0
    mixed_ratio = mixed_count / total_texts if total_texts > 0 else 0

    return {
        'Hebrew Count': hebrew_count,
        'English Count': english_count,
        'Mixed Count': mixed_count,
        'Hebrew Ratio': hebrew_ratio,
        'English Ratio': english_ratio,
        'Mixed Ratio': mixed_ratio
    }


def analyze_and_create_docx(form_data, file_path, output_folder, app):
    # Load the dataset
    try:
        if file_path.endswith('.xlsx'):
            df = pd.read_excel(file_path)
        elif file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
        else:
            messagebox.showerror("Error", "Unsupported file type. Please upload an .xlsx or .csv file.")
            return
    except Exception as e:
        messagebox.showerror("Error", f"Failed to load file: {e}")
        return

    # Extract required fields
    required_columns = ['PatientID', 'Text', 'Record_Date']
    missing_columns = [col for col in required_columns if col not in df.columns]
    if len(missing_columns) > 0:
        messagebox.showerror("Error", f"The following required columns are missing:\n{', '.join(missing_columns)}")
        return
    num_distinct_patients = df['PatientID'].nunique()
    text_lengths = df['Text'].apply(lambda x: len(str(x).split()))
    date_range = pd.to_datetime(df['Record_Date'])

    # Calculate text length statistics
    min_text_length = text_lengths.min()
    avg_text_length = text_lengths.mean()
    max_text_length = text_lengths.max()

    # Date range
    min_date = date_range.min().strftime('%Y-%m-%d')
    max_date = date_range.max().strftime('%Y-%m-%d')

    language_ratio = calculate_language_ratios(df)
    print(language_ratio)
    # Create a new Word document
    doc = Document()

    # header
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "NLP-IHO                                                                                                              2025"
    header_para.style = "Heading 1"

    doc_title = doc.add_heading('Dataset Description', level=1)
    doc_title.runs[0].font.size = Pt(22)

    section_ = doc.add_heading('Dataset Overview', level=2)
    section_.runs[0].font.size = Pt(14)
    para = doc.add_paragraph()
    para.add_run("Organization Name: ").bold = True
    org_name = para.add_run(form_data.get('Organization Name', 'N/A'))
    org_name.font.color.rgb = RGBColor(159, 1, 255)

    para = doc.add_paragraph()
    para.add_run("Number of Records: ").bold = True
    para.add_run(str(len(df)))

    para = doc.add_paragraph()
    para.add_run("Data Domain: ").bold = True
    para.add_run(form_data.get('Data Domain', 'N/A'))

    section_ = doc.add_heading('Cohort Characterization', level=2)
    section_.runs[0].font.size = Pt(14)
    para = doc.add_paragraph()
    para.add_run("Number of Distinct Patients: ").bold = True
    para.add_run(str(num_distinct_patients))

    para = doc.add_paragraph()
    para.add_run("Medical Indications Covered: ").bold = True
    para.add_run(form_data.get('Medical Indications', 'N/A'))

    section_ = doc.add_heading('Text Content Details', level=2)
    section_.runs[0].font.size = Pt(14)
    para = doc.add_paragraph()
    para.add_run("Minimum Text Length (total terms): ").bold = True
    para.add_run(str(min_text_length))

    para = doc.add_paragraph()
    para.add_run("Average Text Length (total terms): ").bold = True
    para.add_run(f"{avg_text_length:.2f}")

    para = doc.add_paragraph()
    para.add_run("Maximum Text Length (total terms): ").bold = True
    para.add_run(str(max_text_length))

    section_ = doc.add_heading('Dates', level=2)
    section_.runs[0].font.size = Pt(14)
    para = doc.add_paragraph()
    para.add_run("Date Range of Records: ").bold = True
    para.add_run(f"{min_date} to {max_date}")

    section_ = doc.add_heading('Data Source', level=2)
    section_.runs[0].font.size = Pt(14)
    para = doc.add_paragraph()
    para.add_run("Original Data Source: ").bold = True
    para.add_run(form_data.get('Original Data Source', 'N/A'))

    section_ = doc.add_heading('Language Analysis', level=2)
    section_.runs[0].font.size = Pt(14)
    para = doc.add_paragraph()
    para.add_run("Hebrew Count: ").bold = True
    para.add_run(str(language_ratio['Hebrew Count']))

    para = doc.add_paragraph()
    para.add_run("English Count: ").bold = True
    para.add_run(str(language_ratio['English Count']))

    para = doc.add_paragraph()
    para.add_run("Mixed Count: ").bold = True
    para.add_run(str(language_ratio['Mixed Count']))

    para = doc.add_paragraph()
    para.add_run("Hebrew Ratio: ").bold = True
    para.add_run(f"{language_ratio['Hebrew Ratio']:.2%}")

    para = doc.add_paragraph()
    para.add_run("English Ratio: ").bold = True
    para.add_run(f"{language_ratio['English Ratio']:.2%}")

    para = doc.add_paragraph()
    para.add_run("Mixed Ratio: ").bold = True
    para.add_run(f"{language_ratio['Mixed Ratio']:.2%}")
    explanation = doc.add_paragraph()
    explanation.add_run("Note: Mixed texts are classified as such when more than 10% of the text words belong to a "
                        "secondary language.").italic = True    # para = doc.add_paragraph()

    # Generate and add graphs
    plt.figure(figsize=(10, 5))
    plt.hist(text_lengths, bins=40, color='blue', alpha=0.7)
    plt.title('Distribution of Text Lengths')
    plt.xlabel('Text Length (Number of Terms)')
    plt.ylabel('Frequency')
    plt.tight_layout()
    text_length_image = io.BytesIO()
    plt.savefig(text_length_image, format='png')
    plt.close()
    text_length_image.seek(0)

    plt.figure(figsize=(10, 5))
    plt.hist(date_range, bins=40, color='green', alpha=0.7)
    plt.title('Distribution of Record Dates')
    plt.xlabel('Date')
    plt.ylabel('Frequency')
    plt.tight_layout()
    date_distribution_image = io.BytesIO()
    plt.savefig(date_distribution_image, format='png')
    plt.close()
    date_distribution_image.seek(0)

    plt.figure(figsize=(10, 5))
    plt.bar(['Hebrew', 'English', 'Mixed'], [
        language_ratio['Hebrew Count'],
        language_ratio['English Count'],
        language_ratio['Mixed Count']
    ], color=['blue', 'green', 'orange'])
    plt.title('Language Distribution')
    plt.xlabel('Language')
    plt.ylabel('Count')
    plt.tight_layout()
    language_distribution_image = io.BytesIO()
    plt.savefig(language_distribution_image, format='png')
    plt.close()
    language_distribution_image.seek(0)

    doc.add_page_break()

    doc.add_paragraph('Distribution of Text Lengths', style='Heading 2')
    doc.add_picture(text_length_image, width=Inches(6))
    doc.add_paragraph('')
    doc.add_paragraph('')  # Add two blank rows

    doc.add_paragraph('Distribution of Record Dates', style='Heading 2')
    doc.add_picture(date_distribution_image, width=Inches(6))

    doc.add_paragraph('')
    doc.add_paragraph('')  # Add two blank rows

    doc.add_paragraph('Language Distribution', style='Heading 2')
    doc.add_picture(language_distribution_image, width=Inches(6))
    explanation = doc.add_paragraph()
    explanation.add_run("Note: Mixed texts are classified as such when more than 10% of the text words belong to a secondary language.").italic = True

    # Save the Word document
    datetime_now = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    output_docx = os.path.join(output_folder, f"data_description_{datetime_now}.docx")
    try:
        doc.save(output_docx)
        messagebox.showinfo("Success", f"Output saved to {output_docx}")
        app.destroy()  # Close the application window
    except Exception as e:
        messagebox.showerror("Error", f"Failed to save output: {e}")


def open_file_dialog(entry_widget):
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx"), ("CSV files", "*.csv")])
    if file_path:
        entry_widget.delete(0, tk.END)
        entry_widget.insert(0, file_path)


def open_folder_dialog(entry_widget):
    folder_path = filedialog.askdirectory()
    if folder_path:
        entry_widget.delete(0, tk.END)
        entry_widget.insert(0, folder_path)


def create_app():
    app = tk.Tk()
    app.title("Dataset Analyzer   |   NLP-IHO")

    # Set the window size
    screen_width = app.winfo_screenwidth()
    screen_height = app.winfo_screenheight()
    app.geometry(f"{int(screen_width / 2)}x{int(screen_height * 0.8)}")

    # Style customization
    app.configure(bg="#f7f9fc")

    # Instructions block
    instructions_frame = tk.Frame(app, bg="#f7f9fc", relief="groove", borderwidth=2)
    instructions_frame.pack(fill="x", padx=20, pady=10)

    tk.Label(instructions_frame, text="Please follow the instructions below:", font=("Arial", 14, "bold"), bg="#f7f9fc",
             fg="#001a66").pack(anchor="w", pady=5)
    instructions = [
        "1. Fill in all fields in Hebrew\\English, clearly and concisely.",
        "2. Select the original data file:",
        "       - Supported formats: Excel (.xlsx) or CSV (.csv).",
        "       - Ensure the file includes the following columns:",
        "               • PatientID: Patient's ID.",
        "               • Record_Date: Date of the record.",
        "               • Text: Free text data.",
        "       Note: The file must follow this structure for successful processing.",
        "3. Choose the output folder where the report (.docx) will be saved.",
        "4. Click Submit to process the data and generate the report."
    ]

    for instruction in instructions:
        tk.Label(instructions_frame, text=instruction, font=("Arial", 11), bg="#f7f9fc", fg="#002080", anchor="w",
                 justify="left").pack(anchor="w")

    def create_field(label_text, entry_variable, tooltip_text, row, placeholder_text, height=1):
        tk.Label(form_frame, text=label_text, bg="#f7f9fc", fg="#2c3e50").grid(row=row, column=0, sticky="w", pady=5)

        def clear_placeholder(event):
            if entry.get("1.0", tk.END).strip() == placeholder_text:
                entry.delete("1.0", tk.END)
                entry.config(fg="black")

        def add_placeholder(event):
            if entry.get("1.0", tk.END).strip() == "":
                entry.insert("1.0", placeholder_text)
                entry.config(fg="gray")

        entry = tk.Text(form_frame, height=height, width=40, fg="gray", wrap="word", relief="solid", borderwidth=1)
        entry.insert("1.0", placeholder_text)
        entry.bind("<FocusIn>", clear_placeholder)
        entry.bind("<FocusOut>", add_placeholder)
        entry.grid(row=row, column=1, sticky="w", pady=5)

        help_button = tk.Button(form_frame, text="?", width=2,
                                command=lambda: messagebox.showinfo(label_text, tooltip_text), bg="#dfe6e9",
                                fg="#2c3e50")
        help_button.grid(row=row, column=2, padx=5)

        return entry

    form_frame = tk.Frame(app, bg="#f7f9fc")
    form_frame.pack(padx=20, pady=10, fill="both", expand=True)

    org_name_var = tk.StringVar()
    domain_var = tk.StringVar()
    source_var = tk.StringVar()

    org_name_entry = create_field("Organization Name:", org_name_var, "Enter the name of the organization.", 0,
                                  "e.g., Soroka, Ichilov...")
    domain_entry = create_field("Data Domain:", domain_var, "What is the data domain? (e.g., imaging, pathology).",
                                1, "e.g., Oncology, Imaging")
    source_entry = create_field("Data Source:", source_var,
                                "What is the source of the data?\n"
                                "For example:\n"
                                "\t 'free text fields, Camillion system'\n"
                                "\t 'PDF reports, Ofek system' ", 2, "e.g., 'PDF reports, Ofek system'")
    indications_entry = create_field("Medical Indications:", None,
                                     "Provide a description of the medical indications for the subpopulation in the dataset.\n"
                                     "For example:\n"
                                     "\t 'women treated for breast cancer between 2020-2021'\n"
                                     "\t 'Helsinki study population(123456789)'\n"
                                     "'\t 'Type2 diabetic patients hospitalized in 2022' ", 3,
                                     "e.g., 'Helsinki study population(123456789)'",
                                     height=5)

    tk.Label(form_frame, text="File to Analyze:", bg="#f7f9fc", fg="#2c3e50").grid(row=5, column=0, sticky="w", pady=5)
    file_path_entry = tk.Entry(form_frame, width=40, relief="solid", borderwidth=1)
    file_path_entry.grid(row=5, column=1, sticky="w", pady=5)
    tk.Button(form_frame, text="Browse", command=lambda: open_file_dialog(file_path_entry), bg="#dfe6e9",
              fg="#2c3e50").grid(row=5, column=2, padx=5)

    tk.Label(form_frame, text="Output Folder:", bg="#f7f9fc", fg="#2c3e50").grid(row=6, column=0, sticky="w", pady=5)
    output_folder_entry = tk.Entry(form_frame, width=40, relief="solid", borderwidth=1)
    output_folder_entry.grid(row=6, column=1, sticky="w", pady=5)
    tk.Button(form_frame, text="Browse", command=lambda: open_folder_dialog(output_folder_entry), bg="#dfe6e9",
              fg="#2c3e50").grid(row=6, column=2, padx=5)

    def on_submit():
        form_data = {
            "Organization Name": org_name_entry.get("1.0", tk.END).strip(),
            "Data Domain": domain_entry.get("1.0", tk.END).strip(),
            "Original Data Source": source_entry.get("1.0", tk.END).strip(),
            "Medical Indications": indications_entry.get("1.0", tk.END).strip(),
            # "Data Collection Method": collection_method_entry.get("1.0", tk.END).strip()
        }
        file_path = file_path_entry.get()
        output_folder = output_folder_entry.get()

        if not all(form_data.values()) or not file_path or not output_folder:
            messagebox.showwarning("Warning", "All fields must be filled.")
            return

        analyze_and_create_docx(form_data, file_path, output_folder, app)

    tk.Button(app, text="Submit", command=on_submit, bg="#0984e3", fg="white", font=("Arial", 12, "bold"),
              relief="flat").pack(pady=15)

    app.mainloop()


if __name__ == "__main__":
    create_app()
